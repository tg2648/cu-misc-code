"""
Converts an Excel sheet into a Word document.

Each column becomes a paragraph where the column header is bold
and the column value follows the header:
    Column 1: column 1 value
    Column 2: column 2 value

Each Excel row becomes split with page breaks in Word
"""

# Standard Library
import os
from pathlib import Path

# Third-party
import docx
import pandas as pd
import numpy as np
from dotenv import load_dotenv

# Load environment variables
basedir = Path(__file__).parent.resolve()
env_path = basedir / '.env'
load_dotenv(dotenv_path=env_path)


SHEET_NAME = 'A&S not decided'
SHEET_PATH = os.getenv('FILE_PATH')

COLUMNS = [
    "Timestamp",
    "Division",
    "Department",
    "Position",
    "Position Description",
    "Candidate",
    "Projected Costs",
    "Funding Sources",
    "Search Status",
    "Grounds for Exception: check all that apply",
    "Joint Hire? Contact",
    "Named Chair",
    "Chair Info",
    "Rationale or Candidate's Qualifications",
]

# HEADING_COLUMN = 'Candidate'  # Column whose value will go into entry headings

df = pd.read_excel(SHEET_PATH, sheet_name=SHEET_NAME, dtype=np.str, keep_default_na=False)
df = df[df['Hiring Exceptions Committee Review'] == 'Yes']

doc = docx.Document()
p = doc.add_paragraph('A&S Exceptions', 'Title')

for i in range(len(df)):

    doc.add_heading(f"Entry {i+1} - {df.iloc[i]['Division']} - {df.iloc[i]['Department']} - {df.iloc[i]['Candidate']}", 1)

    for col in COLUMNS:

        p = doc.add_paragraph()
        r = p.add_run(f'{col}: ')
        r.bold = True
        r = p.add_run(df.iloc[i][col])

    r.add_break(docx.enum.text.WD_BREAK.PAGE)  # Add page break to the last Run

output_name = Path(SHEET_PATH).with_suffix('.docx')
doc.save(output_name)
