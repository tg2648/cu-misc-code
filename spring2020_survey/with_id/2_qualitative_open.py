# Standard Library
import os
from pathlib import Path

# Third-party
import docx

# Local imports
from data import df
from data import qualtrics_columns


def isnotna(text):
    """
    True if text is NOT '' or 'N/A'
    """
    if (text != '') and (text != 'N/A'):
        return True
    else:
        return False


OUTPUT_DIR = os.getenv('OUTPUT_DIR')
FILE_NAME = 'A&S Spring 2020 Survey - Qualitative - with ID.docx'
DOCX_PATH = Path(OUTPUT_DIR) / FILE_NAME

ID_COL = 'ResponseId'
SCHOOL_COLUMN = 'Q2'
schools = sorted(list(df[SCHOOL_COLUMN].unique()))

# Filter questions
question_type_filter = qualtrics_columns['Question_Category'].eq('Qualitative - Open Text')
questions = qualtrics_columns[question_type_filter]

doc = docx.Document(DOCX_PATH)  # Open the same document from part 1
doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)  # Add page break to the last Run
p = doc.add_heading('Responses to open-text questions, by school.', 1)

# Process questions grouped by school - school grouping is based on actual survey data
for school in schools:
    p = doc.add_heading(school, 2)

    # Filter on school
    school_filter = df[SCHOOL_COLUMN].eq(school)
    df_school = df[school_filter]

    for question in questions.itertuples():

        if question.Question_Statement != 'None':
            p = doc.add_heading(f'{question.Question_Text} - {question.Question_Statement}', 3)
        else:
            p = doc.add_heading(f'{question.Question_Text}', 3)

        p.runs[0].add_break()

        # Get all responses for that question
        # Put each response as a paragraph in Word
        responses = df_school[[ID_COL, question.Question_ID]]
        for row in responses.itertuples(index=False):
            if isnotna(getattr(row, question.Question_ID)):  # ignore blank responses
                p = doc.add_paragraph(f'[{getattr(row, ID_COL)}] {getattr(row, question.Question_ID)}')
                p = doc.add_paragraph('---')  # separate each response

doc.save(DOCX_PATH)
