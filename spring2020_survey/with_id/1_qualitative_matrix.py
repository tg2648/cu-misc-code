# Standard Library
import os
from pathlib import Path

# Third-party
import docx

# Local imports
from data import df
from data import qualtrics_columns


def isnotna(text):
    """
    True if text is NOT '' or 'N/A'
    """
    if (text != '') and (text != 'N/A'):
        return True
    else:
        return False


OUTPUT_DIR = os.getenv('OUTPUT_DIR')
FILE_NAME = 'A&S Spring 2020 Survey - Qualitative - with ID.docx'
DOCX_PATH = Path(OUTPUT_DIR) / FILE_NAME
ID_COL = 'ResponseId'

course_types = list(qualtrics_columns['Course_Type'].unique())
course_types.remove('None')
question_type_filter = qualtrics_columns['Question_Category'].eq('Qualitative - Matrix')

doc = docx.Document()
p = doc.add_paragraph('A&S Student Course Survey\nSpring 2020', 'Title')
p = doc.add_paragraph('Excludes responses that answered "Yes" to "Were you in a study abroad program this semester?" (n=129)')
p = doc.add_heading('Qualitative parts of matrix questions, by course type.', 1)

# Process questions grouped by course type
# Course type grouping is based on columns (each course type has dedicated columns in the data)
for course_type in course_types:
    p = doc.add_heading(course_type, 2)

    # Filter on course type and question category
    course_type_filter = qualtrics_columns['Course_Type'].eq(course_type)
    questions = qualtrics_columns[question_type_filter & course_type_filter]

    for question in questions.itertuples():
        # Each qualitative response in matrix type questions has two columns:
        # Q##_##     : Matrix value
        # Q##_##_TEXT: Text entry

        matrix_value_col = question.Question_ID[:-5]
        text_entry_col = question.Question_ID

        p = doc.add_heading(f'{question.Question_Text} - {question.Question_Statement}', 3)
        p.runs[0].add_break()

        # Get all responses for that question
        # Put each response as a paragraph in Word
        responses = df[[ID_COL, matrix_value_col, text_entry_col]].sort_values(by=[matrix_value_col])
        for row in responses.itertuples(index=False):
            if isnotna(getattr(row, text_entry_col)):  # ignore blank responses
                p = doc.add_paragraph()
                r = p.add_run(f'[{getattr(row, ID_COL)}] ')  # use getattr() to index the namedtuple with a variable
                r = p.add_run(f'{getattr(row, matrix_value_col)} - ')
                r.italic = True
                r = p.add_run(getattr(row, text_entry_col))

doc.save(DOCX_PATH)
