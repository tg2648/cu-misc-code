# Standard Library
import os
from pathlib import Path

# Third-party
import docx

# Local imports
from data import df
from data import qualtrics_columns


def isnotna(text):
    """
    True if text is NOT '' or 'N/A'
    """
    if (text != '') and (text != 'N/A'):
        return True
    else:
        return False


OUTPUT_DIR = os.getenv('OUTPUT_DIR')
FILE_NAME = 'A&S Spring 2020 Survey - Qualitative - with ID.docx'
DOCX_PATH = Path(OUTPUT_DIR) / FILE_NAME
ID_COL = 'ResponseId'

course_types = list(qualtrics_columns['Course_Type'].unique())
course_types.remove('None')
question_type_filter = qualtrics_columns['Question_Category'].eq('Qualitative - Course Difficulty')

doc = docx.Document(DOCX_PATH)  # Open the same document from part 1
doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)  # Add page break to the last Run
p = doc.add_heading('Responses to "If you had difficulty with any of the above, please elaborate", by course type.', 1)

# Process questions grouped by course type
# Course type grouping is based on columns (each course type has dedicated columns in the data)
for course_type in course_types:
    p = doc.add_heading(course_type, 2)

    # Filter on course type and question category
    course_type_filter = qualtrics_columns['Course_Type'].eq(course_type)
    questions = qualtrics_columns[question_type_filter & course_type_filter]

    for question in questions.itertuples():

        # Get all responses for that question
        # Put each response as a paragraph in Word
        responses = df[[ID_COL, question.Question_ID]]
        for row in responses.itertuples(index=False):
            if isnotna(getattr(row, question.Question_ID)):  # ignore blank responses
                p = doc.add_paragraph(f'[{getattr(row, ID_COL)}] {getattr(row, question.Question_ID)}')
                p = doc.add_paragraph('---')  # separate each response

doc.save(DOCX_PATH)
